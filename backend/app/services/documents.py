"""One registry for every document the assistant can answer from.

The knowledge base used to be a glob over `data/sops/*.md` with YAML front
matter. That is fine for a corpus you author yourself and useless the moment
somebody uploads a PDF: there is no front matter to read, no `.md` to match,
and no place for the file to live. Any uploaded document would have been a
second, parallel code path — which is exactly how two search systems that
disagree with each other get built.

So documents are a type now, and the SOP corpus is one source of them:

    source="sop"     hand-authored markdown in data/sops/, front matter parsed
    source="upload"  anything a user sent, text extracted per format

Everything downstream — chunking, embedding, retrieval, the viewer, citations
— sees the same shape and does not care which it got. Adding a new format
means adding one extractor here, and nothing else changes.

Uploads are validated before they reach the index. The checks are ordered
cheapest-first, and each one has a reason it exists rather than being defensive
for its own sake:

    extension      the formats we can actually read
    magic bytes    a .pdf that is not a PDF is either a mistake or an attack
    size           a 200MB file blocks the event loop while it embeds
    sanitised name no path separators reach the filesystem
    sha-256        the same manual uploaded twice would double its own weight
                   in retrieval and outrank everything else
    text yield     a scanned PDF extracts to nothing and would index as noise
    relevance      a holiday photo or a payslip in the SOP index makes the ask
                   bar worse for everyone; this is the "filter out irrelevant
                   data" guard
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import re
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from app.config import settings

DOC_DIR = Path(settings.sop_dir).parent / "documents"
ORIGINALS = DOC_DIR / "originals"
REGISTRY = DOC_DIR / "registry.json"

MAX_BYTES = 15 * 1024 * 1024
MIN_EXTRACTED_CHARS = 200
MIN_RELEVANCE_HITS = 3

# Formats we can turn into text. Anything else is refused with its name said
# out loud, so the user knows what to convert to rather than guessing.
ALLOWED = {".md", ".txt", ".pdf", ".docx", ".csv"}

# First bytes that must be present for the formats that have a signature.
_MAGIC = {
    ".pdf": b"%PDF-",
    ".docx": b"PK\x03\x04",  # OOXML is a zip
}


class DocumentRejected(Exception):
    """Refused before indexing. The message is shown to the user verbatim."""


@dataclass
class DocumentMeta:
    doc_id: str
    title: str
    source: str  # "sop" | "upload"
    format: str  # ".pdf", ".md", ...
    department: str = ""
    revision: str = ""
    original_name: str = ""
    stored_name: str = ""
    size_bytes: int = 0
    sha256: str = ""
    uploaded_at: str = ""
    chunks: int = 0
    chars: int = 0

    def as_dict(self) -> dict:
        return asdict(self)


# ---------------------------------------------------------------- extraction


def _from_markdown(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


def _from_csv(raw: bytes) -> str:
    """Render as a markdown table so the chunker sees rows, not one blob.

    Capped: a 10k-row export contributes nothing to retrieval past the first
    page, and would swamp the index with near-identical chunks.
    """
    text = raw.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ""
    header, body = rows[0], rows[1:200]
    lines = ["# Data", "", "| " + " | ".join(header) + " |",
             "| " + " | ".join("---" for _ in header) + " |"]
    lines += ["| " + " | ".join(cell.replace("|", "/") for cell in row) + " |" for row in body]
    if len(rows) > 201:
        lines.append(f"\n_{len(rows) - 201} further rows not shown._")
    return "\n".join(lines)


def _from_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        # A page heading gives the chunker a boundary and gives citations
        # something to name, since a PDF has no sections to speak of.
        parts.append(f"# Page {number}\n\n{text}")
    return "\n\n".join(parts)


def _from_docx(raw: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw))
    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 3)} {text}")
        elif style.startswith("title"):
            parts.append(f"# {text}")
        else:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("|", "/") for c in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(parts)


_EXTRACTORS = {
    ".md": _from_markdown,
    ".txt": _from_markdown,
    ".csv": _from_csv,
    ".pdf": _from_pdf,
    ".docx": _from_docx,
}


def extract_text(suffix: str, raw: bytes) -> str:
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        raise DocumentRejected(f"{suffix} files are not supported.")
    try:
        return extractor(raw).strip()
    except DocumentRejected:
        raise
    except Exception as exc:  # a corrupt file should not 500 the endpoint
        raise DocumentRejected(f"Could not read this {suffix} file: {exc}") from exc


# ---------------------------------------------------------------- relevance


def _vocabulary() -> set[str]:
    """Manufacturing terms, built from the plant's own data where possible.

    Deriving most of it from machines.json means a plant that runs different
    equipment gets a different filter without anyone editing this file.
    """
    terms = {
        "sop", "procedure", "operator", "maintenance", "machine", "line",
        "shift", "quality", "defect", "inspection", "tolerance", "calibration",
        "changeover", "downtime", "safety", "lockout", "torque", "assembly",
        "production", "batch", "work order", "scrap", "rework", "audit",
        "preventive", "spindle", "fixture", "tooling", "material", "supplier",
        "iso", "revision", "hazard", "ppe", "gauge", "conveyor",
    }
    try:
        from app.services.data_loader import load

        machines = load()["machines"]
        for row in machines.itertuples():
            terms.add(str(row.machine_type).lower())
            terms.add(str(row.line).lower())
            for keyword in (row.keywords if isinstance(row.keywords, list) else []):
                terms.add(str(keyword).lower())
    except Exception:
        pass  # a filter built from the static list alone is still a filter
    return terms


_vocab: set[str] | None = None


def relevance_hits(text: str) -> set[str]:
    global _vocab
    if _vocab is None:
        _vocab = _vocabulary()
    lowered = text.lower()
    return {term for term in _vocab if term in lowered}


# ---------------------------------------------------------------- registry


def _load_registry() -> dict[str, dict]:
    if not REGISTRY.is_file():
        return {}
    try:
        return json.loads(REGISTRY.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        print("[documents] registry.json unreadable, starting a fresh one")
        return {}


def _save_registry(registry: dict[str, dict]) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    REGISTRY.write_text(json.dumps(registry, indent=2), encoding="utf-8")


def _safe_stem(name: str) -> str:
    """Filename to something that cannot escape the uploads directory."""
    stem = Path(name).name  # drops any directory component, including "..\"
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "_", stem).strip("._ ")
    return stem[:80] or "document"


def _next_doc_id(registry: dict[str, dict]) -> str:
    used = {
        int(m.group(1))
        for key in registry
        if (m := re.fullmatch(r"DOC-(\d{3})", key))
    }
    n = 1
    while n in used:
        n += 1
    return f"DOC-{n:03d}"


# ---------------------------------------------------------------- SOP corpus


def _sop_documents() -> list[DocumentMeta]:
    """The hand-authored corpus, presented as documents like any other."""
    from app.services.knowledge_base import KnowledgeBase

    out: list[DocumentMeta] = []
    for path in sorted(Path(settings.sop_dir).glob("*.md")):
        raw = path.read_text(encoding="utf-8")
        meta, body = KnowledgeBase._parse_front_matter(raw)
        out.append(
            DocumentMeta(
                doc_id=meta.get("doc_id", path.stem),
                title=meta.get("title", path.stem),
                source="sop",
                format=".md",
                department=meta.get("department", ""),
                revision=meta.get("revision", ""),
                original_name=path.name,
                stored_name=str(path),
                size_bytes=path.stat().st_size,
                uploaded_at=datetime.fromtimestamp(
                    path.stat().st_mtime, tz=timezone.utc
                ).isoformat(timespec="seconds"),
                chars=len(body),
            )
        )
    return out


# ---------------------------------------------------------------- public API


def list_documents() -> list[DocumentMeta]:
    """Everything answerable, SOPs and uploads together, newest upload first."""
    uploads = [DocumentMeta(**row) for row in _load_registry().values()]
    uploads.sort(key=lambda d: d.uploaded_at, reverse=True)
    return _sop_documents() + uploads


def get_meta(doc_id: str) -> DocumentMeta | None:
    wanted = doc_id.strip().upper()
    for doc in list_documents():
        if doc.doc_id.upper() == wanted:
            return doc
    return None


def get_text(doc_id: str) -> str | None:
    """Readable text for the viewer, whatever the file started as."""
    meta = get_meta(doc_id)
    if meta is None:
        return None

    path = Path(meta.stored_name)
    if not path.is_file():
        return None

    if meta.source == "sop":
        from app.services.knowledge_base import KnowledgeBase

        _, body = KnowledgeBase._parse_front_matter(path.read_text(encoding="utf-8"))
        return body
    return extract_text(meta.format, path.read_bytes())


def original_path(doc_id: str) -> Path | None:
    """The file exactly as uploaded, for download."""
    meta = get_meta(doc_id)
    if meta is None:
        return None
    path = Path(meta.stored_name)
    return path if path.is_file() else None


def save_upload(filename: str, raw: bytes, department: str = "") -> DocumentMeta:
    """Validate, store, extract and register. Raises DocumentRejected."""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED:
        raise DocumentRejected(
            f"{suffix or 'This file type'} is not supported. "
            f"Accepted: {', '.join(sorted(ALLOWED))}."
        )
    if not raw:
        raise DocumentRejected("The file is empty.")
    if len(raw) > MAX_BYTES:
        raise DocumentRejected(
            f"File is {len(raw) / 1_048_576:.1f} MB. The limit is "
            f"{MAX_BYTES // 1_048_576} MB."
        )

    signature = _MAGIC.get(suffix)
    if signature and not raw.startswith(signature):
        raise DocumentRejected(
            f"This does not look like a real {suffix} file — its contents do not "
            "match its extension."
        )

    digest = hashlib.sha256(raw).hexdigest()
    registry = _load_registry()
    for existing in registry.values():
        if existing.get("sha256") == digest:
            raise DocumentRejected(
                f"Already uploaded as {existing['doc_id']} "
                f"({existing['original_name']})."
            )

    text = extract_text(suffix, raw)
    if len(text) < MIN_EXTRACTED_CHARS:
        raise DocumentRejected(
            "No readable text could be extracted. A scanned PDF needs running "
            "through OCR before it can be searched."
        )

    hits = relevance_hits(text)
    if len(hits) < MIN_RELEVANCE_HITS:
        raise DocumentRejected(
            "This does not read like manufacturing documentation, so it was not "
            "added — indexing it would make search worse for everyone. Upload "
            "SOPs, manuals, work instructions or audit records."
        )

    doc_id = _next_doc_id(registry)
    stem = _safe_stem(filename)
    ORIGINALS.mkdir(parents=True, exist_ok=True)
    stored = ORIGINALS / f"{doc_id}__{stem}"
    stored.write_bytes(raw)

    # First markdown heading, else the filename. A document with no title at
    # all is unciteable in the ask bar.
    heading = re.search(r"^#{1,3}\s+(.+)$", text, re.MULTILINE)
    title = (heading.group(1).strip() if heading else Path(stem).stem.replace("_", " "))[:120]

    meta = DocumentMeta(
        doc_id=doc_id,
        title=title,
        source="upload",
        format=suffix,
        department=department.strip()[:60],
        revision="",
        original_name=Path(filename).name,
        stored_name=str(stored),
        size_bytes=len(raw),
        sha256=digest,
        uploaded_at=datetime.now(timezone.utc).isoformat(timespec="seconds"),
        chars=len(text),
    )

    # Index before registering: a document in the registry that is not in the
    # index would be listed as searchable and silently never match.
    from app.services.knowledge_base import get_knowledge_base

    meta.chunks = get_knowledge_base().index_document(
        doc_id=meta.doc_id,
        title=meta.title,
        text=text,
        department=meta.department,
        revision=meta.revision,
        source="upload",
    )

    registry[doc_id] = meta.as_dict()
    _save_registry(registry)
    print(f"[documents] {doc_id} indexed, {meta.chunks} chunks, matched {len(hits)} terms")
    return meta


def delete_upload(doc_id: str) -> bool:
    """Remove an upload and its chunks. SOPs are not deletable from the UI."""
    registry = _load_registry()
    row = registry.get(doc_id.upper())
    if row is None:
        return False

    from app.services.knowledge_base import get_knowledge_base

    get_knowledge_base().remove_document(doc_id.upper())

    path = Path(row["stored_name"])
    if path.is_file():
        path.unlink()
    registry.pop(doc_id.upper())
    _save_registry(registry)
    return True


def indexable() -> Iterable[dict]:
    """Every document that should be in the vector index, for a full rebuild."""
    for meta in list_documents():
        text = get_text(meta.doc_id)
        if not text:
            continue
        yield {
            "doc_id": meta.doc_id,
            "title": meta.title,
            "text": text,
            "department": meta.department,
            "revision": meta.revision,
            "source": meta.source,
        }
